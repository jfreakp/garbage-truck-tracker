from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Color palette ─────────────────────────────────────────────────────────────
GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)   # #1B5E20
GREEN_MID   = RGBColor(0x2E, 0x7D, 0x32)   # #2E7D32
GREEN_LIGHT = RGBColor(0x4C, 0xAF, 0x50)   # #4CAF50
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT   = RGBColor(0x42, 0x42, 0x42)
GRAY_ROW    = RGBColor(0xF1, 0xF8, 0xE9)   # very light green tint

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: bold + colored run ────────────────────────────────────────────────
def add_run(para, text, bold=False, color=None, size=None, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    return run

# ── Helper: heading paragraph ─────────────────────────────────────────────────
def section_heading(doc, text, level=1):
    """Green-shaded section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    if level == 1:
        shade_paragraph(p, '1B5E20')
        run = p.add_run(f'  {text}')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = WHITE
    else:
        shade_paragraph(p, '2E7D32')
        run = p.add_run(f'  {text}')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_before = Pt(60)
cover_title.paragraph_format.space_after  = Pt(6)
shade_paragraph(cover_title, '1B5E20')
r = cover_title.add_run('  MEMORIA TÉCNICA DEL SISTEMA  ')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = WHITE

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub1, '2E7D32')
r2 = sub1.add_run('  EcoRuta – Sistema de Rastreo de Recolección de Residuos  ')
r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = WHITE

doc.add_paragraph()

meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.style = 'Table Grid'
meta_data = [
    ('Versión del documento', '1.0'),
    ('Fecha de elaboración',  datetime.date.today().strftime('%d de %B de %Y')),
    ('Plataformas objetivo',  'Web · Android · iOS'),
    ('Presupuesto estimado',  'USD $10,000'),
    ('Clasificación',         'Uso interno / Licitación'),
]
col_w = [Cm(5.5), Cm(9)]
for i, (label, value) in enumerate(meta_data):
    row = meta_table.rows[i]
    row.cells[0].width = col_w[0]
    row.cells[1].width = col_w[1]
    shade_cell(row.cells[0], '2E7D32')
    shade_cell(row.cells[1], 'F1F8E9')
    lp = row.cells[0].paragraphs[0]
    lp.add_run(label).bold = True
    lp.runs[0].font.color.rgb = WHITE
    lp.runs[0].font.size = Pt(10)
    vp = row.cells[1].paragraphs[0]
    vp.add_run(value).font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '1. INTRODUCCIÓN', 1)
intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(8)
add_run(intro,
    'EcoRuta es una plataforma digital multiplataforma diseñada para modernizar el servicio '
    'de recolección de residuos sólidos urbanos. Permite a la ciudadanía conocer en tiempo real '
    'la ubicación del camión recolector, recibir alertas anticipadas cuando el vehículo se '
    'aproxima a su sector y saber qué tipo de residuo —orgánico (negro) o reciclable (verde)— '
    'debe sacar ese día. El sistema se despliega simultáneamente como aplicación web, aplicación '
    'Android y aplicación iOS, garantizando cobertura universal independientemente del dispositivo '
    'del ciudadano.',
    color=GRAY_TEXT, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  2. OBJETIVOS
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '2. OBJETIVOS DEL SISTEMA', 1)

section_heading(doc, '2.1 Objetivo General', 2)
p = doc.add_paragraph()
add_run(p,
    'Desarrollar e implementar una solución tecnológica integral que mejore la eficiencia operativa '
    'del servicio de recolección de residuos y eleve la calidad de vida de la ciudadanía mediante '
    'información oportuna y accesible.',
    color=GRAY_TEXT, size=11)

section_heading(doc, '2.2 Objetivos Específicos', 2)
specifics = [
    'Notificar a los ciudadanos en tiempo real cuando el camión recolector se encuentre a menos de 500 m de su domicilio.',
    'Informar diariamente el tipo de residuo a disponer (tacho verde – reciclables / tacho negro – orgánicos).',
    'Visualizar en un mapa interactivo la posición GPS del recolector con actualización cada 5 segundos.',
    'Gestionar múltiples rutas y vehículos desde un panel de administración centralizado.',
    'Reducir quejas ciudadanas por falta de información sobre el servicio de recolección.',
]
for item in specifics:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    add_run(p, item, color=GRAY_TEXT, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  3. ALCANCE
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '3. ALCANCE', 1)
p = doc.add_paragraph()
add_run(p,
    'El presente documento abarca el diseño, desarrollo, pruebas e implementación de los siguientes '
    'módulos y plataformas:',
    color=GRAY_TEXT, size=11)

alcance_items = [
    ('Web App', 'Dashboard administrativo y portal ciudadano accesible desde cualquier navegador moderno.'),
    ('Android App', 'Aplicación nativa compilada como APK para distribución en Google Play Store o instalación directa.'),
    ('iOS App', 'Aplicación nativa para iPhone/iPad distribuida a través de la App Store de Apple.'),
    ('Backend / API', 'Servidor REST con Next.js 16, base de datos PostgreSQL con PostGIS y notificaciones push vía Firebase FCM.'),
    ('Infraestructura', 'Despliegue en nube con CI/CD, monitoreo y escalabilidad horizontal.'),
]
for title, desc in alcance_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, f'• {title}: ', bold=True, color=GREEN_MID, size=11)
    add_run(p, desc, color=GRAY_TEXT, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  4. FUNCIONALIDADES CLAVE (BONDADES DEL SISTEMA)
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '4. FUNCIONALIDADES CLAVE DEL SISTEMA', 1)

features = [
    ('🔔 Notificación de Proximidad del Recolector',
     'Alertas push enviadas automáticamente a los dispositivos móviles de los ciudadanos cuando el '
     'camión recolector se encuentra a menos de 500 metros de su domicilio. El ciudadano dispone de '
     'tiempo suficiente para sacar su basura antes de que el vehículo llegue. Las notificaciones '
     'funcionan incluso con la aplicación cerrada gracias a Firebase Cloud Messaging (FCM).'),
    ('♻️ Aviso del Tipo de Tacho (Verde / Negro)',
     'El sistema calcula diariamente, según el calendario de recolección configurado, qué tipo de '
     'residuo le corresponde al ciudadano ese día:\n'
     '   • Tacho VERDE → Residuos reciclables (plástico, papel, vidrio, metales).\n'
     '   • Tacho NEGRO → Residuos orgánicos e indiferenciados.\n'
     'La alerta se envía la noche anterior y se muestra como widget destacado en la app.'),
    ('📍 Rastreo GPS en Tiempo Real',
     'Mapa interactivo (Leaflet / OpenStreetMap) que muestra la posición exacta del camión '
     'recolector con actualización automática cada 5 segundos. El ciudadano puede ver la ruta '
     'recorrida, la ruta planificada y el tiempo estimado de llegada a su cuadra. Los operadores '
     'municipales visualizan todos los vehículos de la flota simultáneamente desde el panel de '
     'administración.'),
    ('🗺️ Panel de Administración Municipal',
     'Interfaz exclusiva para funcionarios con control de acceso por roles (ADMIN / USER). Permite '
     'gestionar rutas, horarios, calendarios de recolección diferenciada y ver reportes de '
     'cumplimiento de servicio.'),
    ('📲 Multiplataforma (Web + Android + iOS)',
     'Una única experiencia de usuario coherente disponible en las tres plataformas más utilizadas, '
     'maximizando el alcance a toda la ciudadanía sin importar su dispositivo.'),
]

for i, (title, body) in enumerate(features):
    shade = 'F1F8E9' if i % 2 == 0 else 'E8F5E9'
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after  = Pt(2)
    shade_paragraph(p_title, '4CAF50')
    add_run(p_title, f'  {title}', bold=True, color=WHITE, size=11)

    for line in body.split('\n'):
        p_body = doc.add_paragraph()
        p_body.paragraph_format.left_indent  = Cm(0.5)
        p_body.paragraph_format.space_after  = Pt(2)
        shade_paragraph(p_body, shade)
        add_run(p_body, f'  {line}', color=GRAY_TEXT, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  5. ARQUITECTURA TÉCNICA
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '5. ARQUITECTURA TÉCNICA', 1)

arch_items = [
    ('Frontend Web', 'Next.js 16, React 19, Tailwind CSS 4, Leaflet / React-Leaflet'),
    ('Aplicaciones Móviles', 'React Native (Expo) – código compartido ~80 % entre Android e iOS'),
    ('Backend / API', 'Next.js API Routes (REST), validación Zod, autenticación JWT (jose)'),
    ('Base de Datos', 'PostgreSQL 17 + extensión PostGIS (geometrías espaciales)'),
    ('ORM', 'Prisma 7 con adaptador pg para consultas nativas PostGIS'),
    ('Notificaciones Push', 'Firebase Admin SDK → FCM (Android + iOS + Web)'),
    ('Autenticación', 'JWT httpOnly cookies con hashing bcrypt; roles ADMIN / USER'),
    ('Infraestructura', 'Vercel (Web), Expo EAS Build (Móvil), Supabase / Railway (DB)'),
    ('CI/CD', 'GitHub Actions: lint → test → build → deploy automático en cada PR'),
]

arch_table = doc.add_table(rows=len(arch_items)+1, cols=2)
arch_table.style = 'Table Grid'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = arch_table.rows[0]
for cell, text in zip(hdr.cells, ['Capa / Componente', 'Tecnología']):
    shade_cell(cell, '1B5E20')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=True, color=WHITE, size=10)

for i, (layer, tech) in enumerate(arch_items):
    row = arch_table.rows[i+1]
    shade = 'F1F8E9' if i % 2 == 0 else 'FFFFFF'
    shade_cell(row.cells[0], '2E7D32')
    shade_cell(row.cells[1], shade)
    add_run(row.cells[0].paragraphs[0], layer, bold=True, color=WHITE, size=10)
    add_run(row.cells[1].paragraphs[0], tech, color=GRAY_TEXT, size=10)

# ══════════════════════════════════════════════════════════════════════════════
#  6. DIAGRAMA DE FLUJO (Narrativo)
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '6. FLUJO DE OPERACIÓN', 1)
flow_steps = [
    ('Paso 1 – Registro de ciudadano',
     'El usuario descarga la app (iOS/Android) o accede a la web, se registra con nombre, '
     'correo y dirección. El sistema geocodifica la dirección y almacena las coordenadas.'),
    ('Paso 2 – Configuración de calendario',
     'El administrador municipal carga el calendario de recolección: fechas, rutas, tipo de '
     'residuo y horarios estimados por sector.'),
    ('Paso 3 – Inicio de ruta',
     'El operador del camión activa la sesión de rastreo en la app; el dispositivo comienza a '
     'enviar coordenadas GPS cada 5 segundos al servidor.'),
    ('Paso 4 – Rastreo en tiempo real',
     'Todos los ciudadanos de la ciudad pueden ver en el mapa la posición del camión '
     'actualizada en vivo.'),
    ('Paso 5 – Disparo de notificación de proximidad',
     'Cuando el servidor calcula que el camión está ≤ 500 m de un domicilio registrado, '
     'envía automáticamente una notificación push vía FCM al dispositivo del ciudadano.'),
    ('Paso 6 – Alerta de tipo de residuo',
     'La noche anterior al día de recolección, el sistema envía una notificación indicando '
     '"Mañana: tacho VERDE" o "Mañana: tacho NEGRO" según el calendario configurado.'),
    ('Paso 7 – Cierre de ruta',
     'Al finalizar, el operador cierra la sesión; el sistema registra el recorrido completado '
     'y genera un reporte automático de cumplimiento.'),
]
for step, desc in flow_steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    add_run(p, f'{step}: ', bold=True, color=GREEN_MID, size=11)
    add_run(p, desc, color=GRAY_TEXT, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  7. REQUERIMIENTOS NO FUNCIONALES
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '7. REQUERIMIENTOS NO FUNCIONALES', 1)
non_func = [
    ('Disponibilidad',       '99.5 % uptime mensual (SLA)'),
    ('Latencia GPS',         '< 5 s de retardo en actualización de posición'),
    ('Notificaciones',       '< 3 s desde el disparo del evento hasta la entrega push'),
    ('Seguridad',            'Comunicaciones HTTPS/TLS 1.3; JWT con expiración 1 h; bcrypt cost 12'),
    ('Escalabilidad',        'Soporta hasta 50,000 usuarios concurrentes sin degradación'),
    ('Compatibilidad móvil', 'Android 10+ / iOS 15+'),
    ('Compatibilidad web',   'Chrome 120+, Firefox 120+, Safari 17+, Edge 120+'),
    ('Accesibilidad',        'WCAG 2.1 nivel AA en interfaz web'),
]
nf_table = doc.add_table(rows=len(non_func)+1, cols=2)
nf_table.style = 'Table Grid'
hdr2 = nf_table.rows[0]
for cell, text in zip(hdr2.cells, ['Atributo', 'Especificación']):
    shade_cell(cell, '2E7D32')
    hdr2.cells[0].width = Cm(4.5)
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, text, bold=True, color=WHITE, size=10)
for i, (attr, spec) in enumerate(non_func):
    row = nf_table.rows[i+1]
    shade = 'F1F8E9' if i % 2 == 0 else 'FFFFFF'
    shade_cell(row.cells[0], '4CAF50')
    shade_cell(row.cells[1], shade)
    add_run(row.cells[0].paragraphs[0], attr, bold=True, color=WHITE, size=10)
    add_run(row.cells[1].paragraphs[0], spec, color=GRAY_TEXT, size=10)

# ══════════════════════════════════════════════════════════════════════════════
#  8. PLAN DE IMPLEMENTACIÓN
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '8. PLAN DE IMPLEMENTACIÓN', 1)
phases = [
    ('Fase 1 – Análisis y Diseño (Semanas 1-2)',
     ['Levantamiento de requerimientos con el municipio',
      'Diseño de arquitectura y base de datos',
      'Wireframes y prototipo de UI/UX (Figma)',
      'Configuración de repositorio y entornos CI/CD']),
    ('Fase 2 – Backend y API (Semanas 3-5)',
     ['Implementación de modelos Prisma + PostGIS',
      'API REST: autenticación, rutas, vehículos, calendarios',
      'Lógica de detección de proximidad y disparo FCM',
      'Tests de integración y documentación OpenAPI']),
    ('Fase 3 – Web App (Semanas 4-6)',
     ['Portal ciudadano: mapa en tiempo real, widget de residuo',
      'Panel de administración: CRUD rutas, calendarios, usuarios',
      'Responsive design para móvil/tablet/escritorio']),
    ('Fase 4 – Apps Móviles (Semanas 5-8)',
     ['Desarrollo React Native compartido (Android + iOS)',
      'Integración notificaciones push FCM/APNs',
      'Geolocalización en segundo plano',
      'Publicación en Google Play Store y App Store']),
    ('Fase 5 – QA, Piloto y Lanzamiento (Semanas 9-10)',
     ['Pruebas de carga y seguridad',
      'Piloto con grupo de 200 ciudadanos',
      'Corrección de bugs y ajustes de UX',
      'Lanzamiento oficial y capacitación municipal']),
]
for phase_title, tasks in phases:
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(8)
    shade_paragraph(ph, '4CAF50')
    add_run(ph, f'  {phase_title}', bold=True, color=WHITE, size=11)
    for task in tasks:
        tp = doc.add_paragraph(style='List Bullet')
        tp.paragraph_format.left_indent = Cm(1.2)
        add_run(tp, task, color=GRAY_TEXT, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  9. PRESUPUESTO
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '9. PRESUPUESTO ESTIMADO', 1)

p_intro = doc.add_paragraph()
add_run(p_intro,
    'El presupuesto total del proyecto asciende a ',
    color=GRAY_TEXT, size=11)
add_run(p_intro, 'USD $10,000', bold=True, color=GREEN_DARK, size=12)
add_run(p_intro,
    ', distribuido en los rubros que se detallan a continuación:',
    color=GRAY_TEXT, size=11)

budget_items = [
    ('Análisis, Diseño y Arquitectura',        '10 %',  '$1,000',
     'Levantamiento de reqs, diseño de BD, wireframes UX'),
    ('Desarrollo Backend + API',               '20 %',  '$2,000',
     'Next.js API, Prisma/PostGIS, FCM, autenticación JWT'),
    ('Desarrollo Web App (Portal + Admin)',     '20 %',  '$2,000',
     'Next.js frontend, mapa Leaflet, dashboard admin RBAC'),
    ('Desarrollo Android App',                 '15 %',  '$1,500',
     'React Native (Expo), notificaciones push, GPS'),
    ('Desarrollo iOS App',                     '15 %',  '$1,500',
     'React Native (Expo), notificaciones APNs, GPS'),
    ('QA, Pruebas y Seguridad',                 '8 %',   '$800',
     'Pruebas de carga, penetración, compatibilidad'),
    ('Infraestructura y Despliegue (1 año)',    '7 %',   '$700',
     'Vercel, Supabase/Railway, dominio, certificados SSL'),
    ('Documentación y Capacitación',            '5 %',   '$500',
     'Manuales de usuario y admin, sesión de capacitación'),
    ('TOTAL',                                 '100 %', '$10,000', ''),
]

bud_table = doc.add_table(rows=len(budget_items)+1, cols=4)
bud_table.style = 'Table Grid'
bud_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdrs = ['Rubro', 'Porcentaje', 'Monto (USD)', 'Descripción']
for cell, text in zip(bud_table.rows[0].cells, hdrs):
    shade_cell(cell, '1B5E20')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=True, color=WHITE, size=10)

for i, (rubro, pct, monto, desc) in enumerate(budget_items):
    row = bud_table.rows[i+1]
    is_total = rubro == 'TOTAL'
    shade = '2E7D32' if is_total else ('F1F8E9' if i % 2 == 0 else 'FFFFFF')
    txt_color = WHITE if is_total else GRAY_TEXT
    for cell in row.cells:
        shade_cell(cell, shade)
    for cell, text in zip(row.cells, [rubro, pct, monto, desc]):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text != desc else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, bold=is_total, color=txt_color, size=10)

doc.add_paragraph()
note = doc.add_paragraph()
note.paragraph_format.left_indent = Cm(0.5)
add_run(note,
    'Nota: ',
    bold=True, color=GREEN_MID, size=10)
add_run(note,
    'El presupuesto cubre el desarrollo completo hasta el lanzamiento oficial. '
    'El mantenimiento correctivo post-lanzamiento (6 meses) está incluido sin costo adicional. '
    'Costos de publicación en App Store (USD $99/año) y Google Play (USD $25 pago único) '
    'no están incluidos y corren por cuenta del contratante.',
    color=GRAY_TEXT, size=10)

# ══════════════════════════════════════════════════════════════════════════════
#  10. EQUIPO DE TRABAJO
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '10. EQUIPO DE TRABAJO', 1)
team = [
    ('1', 'Líder de Proyecto / Arquitecto', '10 semanas', '20 %'),
    ('1', 'Desarrollador Full-Stack Senior (Web + API)', '10 semanas', '30 %'),
    ('1', 'Desarrollador Mobile (React Native)', '8 semanas',  '25 %'),
    ('1', 'Diseñador UI/UX',                     '4 semanas',  '10 %'),
    ('1', 'Especialista QA / Tester',             '3 semanas',  '10 %'),
    ('1', 'DevOps / Infraestructura',             '2 semanas',   '5 %'),
]
team_table = doc.add_table(rows=len(team)+1, cols=4)
team_table.style = 'Table Grid'
for cell, text in zip(team_table.rows[0].cells, ['#', 'Rol', 'Duración', 'Dedicación']):
    shade_cell(cell, '2E7D32')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=True, color=WHITE, size=10)
for i, row_data in enumerate(team):
    row = team_table.rows[i+1]
    shade = 'F1F8E9' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        shade_cell(cell, shade)
    for cell, text in zip(row.cells, row_data):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cell.paragraphs[0], text, color=GRAY_TEXT, size=10)

# ══════════════════════════════════════════════════════════════════════════════
#  11. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '11. CONCLUSIONES Y BENEFICIOS ESPERADOS', 1)
conclusions = [
    'Reducción del 60–70 % de quejas ciudadanas por falta de información sobre el servicio.',
    'Incremento de la tasa de disposición correcta de residuos diferenciados (verde/negro) gracias a las alertas preventivas.',
    'Optimización operativa para el municipio: rutas más eficientes y control de cumplimiento en tiempo real.',
    'Plataforma escalable que puede extenderse a otros servicios municipales (agua, alumbrado, etc.).',
    'Inversión inicial de USD $10,000 con retorno tangible en satisfacción ciudadana, reducción de reclamos y eficiencia operativa en el primer año.',
]
for c in conclusions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(5)
    add_run(p, '✔  ', bold=True, color=GREEN_LIGHT, size=11)
    add_run(p, c, color=GRAY_TEXT, size=11)

# Footer note
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(footer_p, '1B5E20')
add_run(footer_p,
    '  EcoRuta © 2026 — Documento confidencial para uso interno y procesos de licitación  ',
    color=WHITE, size=9)

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\jfreakp\Documents\work\garbage-truck-tracker\EcoRuta_Memoria_Tecnica.docx'
doc.save(out_path)
print(f'Documento generado: {out_path}')
